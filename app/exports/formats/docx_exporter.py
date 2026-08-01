"""Microsoft Word (.docx) export via python-docx, preserving paragraph and
page structure.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from app.core.models import OCRDocument


def export(document: OCRDocument, path: str | Path) -> Path:
    path = Path(path)
    docx = DocxDocument()

    title = docx.add_heading(Path(document.source_path).name, level=1)
    for run in title.runs:
        run.font.size = Pt(16)

    meta = docx.add_paragraph()
    meta.add_run(
        f"Engine: {document.engine}  |  Language: {document.language}  |  "
        f"Confidence: {document.mean_confidence:.1f}%"
    ).italic = True

    for page in document.pages:
        if document.page_count > 1:
            docx.add_heading(f"Page {page.page_number}", level=2)
        for paragraph in page.text.split("\n\n"):
            if paragraph.strip():
                docx.add_paragraph(paragraph.strip())
        if page.page_number != document.pages[-1].page_number:
            docx.add_page_break()

    docx.save(str(path))
    return path
